from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "reports"
BASE_REPORT = ROOT / "final_report.md"
COMPLETE_MD = REPORTS_DIR / "final_report_complete.md"
DOCX_OUT = REPORTS_DIR / "final_report_complete.docx"
PDF_OUT = REPORTS_DIR / "final_report_complete.pdf"

CODE_APPENDICES = [
    ("automation/main.py", ROOT / "automation" / "main.py"),
    ("automation/translator.py", ROOT / "automation" / "translator.py"),
    ("automation/content_enricher.py", ROOT / "automation" / "content_enricher.py"),
    ("automation/image_generator.py", ROOT / "automation" / "image_generator.py"),
    ("automation/strapi_api.py", ROOT / "automation" / "strapi_api.py"),
    ("automation/data/places.json", ROOT / "automation" / "data" / "places.json"),
]


def build_complete_markdown() -> str:
    report = BASE_REPORT.read_text(encoding="utf-8").rstrip()
    parts = [report, "\n\n---\n\n# Ekler: Python Betiğinin Tam Metni\n"]
    for label, path in CODE_APPENDICES:
        language = "json" if path.suffix == ".json" else "python"
        code = path.read_text(encoding="utf-8")
        parts.append(f"\n## Ek - {label}\n\n```{language}\n{code.rstrip()}\n```\n")
    return "\n".join(parts).strip() + "\n"


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def parse_markdown_lines(markdown_text: str):
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            lang = line.strip("`").strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            yield ("code", lang, "\n".join(code_lines))
        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[i + 1]):
            header = split_table_row(line)
            rows = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            yield ("table", header, rows)
            continue
        else:
            yield ("line", line)
        i += 1


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def add_docx_text(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(clean_inline_markdown(text))
    return paragraph


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths_dxa: list[int]):
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side in ("top", "bottom", "start", "end"):
        element = tbl_cell_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), "80" if side in {"top", "bottom"} else "120")
        element.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in col_widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths_dxa[min(idx, len(col_widths_dxa) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_run_font(run, name: str):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)


def configure_docx_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]
    for style_name, size, color, before, after in heading_tokens:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    code_style = styles["CodeBlock"] if "CodeBlock" in styles else styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(7)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.left_indent = Inches(0.08)
    code_style.paragraph_format.right_indent = Inches(0.08)
    code_style.paragraph_format.line_spacing = 1.0

    callout_style = styles["Callout"] if "Callout" in styles else styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout_style.font.name = "Calibri"
    callout_style.font.size = Pt(10)
    callout_style.font.italic = True
    callout_style.font.color.rgb = RGBColor.from_string("1F3A5F")
    callout_style.paragraph_format.space_before = Pt(4)
    callout_style.paragraph_format.space_after = Pt(6)


def setup_docx_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Dünyayi Gezayisun AI Rehberuylan"
    for run in footer.runs:
        set_run_font(run, "Calibri")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")


def style_screenshot_placeholder(paragraph):
    if "Eklenmesi gereken ekran görüntüsü:" not in paragraph.text:
        return
    paragraph.style = "Callout"
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)


def add_code_block(doc: Document, code: str):
    paragraph = doc.add_paragraph(style="CodeBlock")
    run = paragraph.add_run(code)
    set_run_font(run, "Consolas")
    run.font.size = Pt(6.6 if len(code.splitlines()) > 80 else 7.4)
    paragraph.paragraph_format.keep_together = False
    return paragraph


def add_formatted_table(doc: Document, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"

    col_width = 9360 // max(1, len(header))
    col_widths = [col_width] * len(header)
    col_widths[-1] += 9360 - sum(col_widths)

    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = clean_inline_markdown(header[idx])
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, "Calibri")
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor.from_string("0B2545")

    for row in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            cell.text = clean_inline_markdown(row[idx] if idx < len(row) else "")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    set_run_font(run, "Calibri")
                    run.font.size = Pt(10)

    set_table_geometry(table, col_widths)
    doc.add_paragraph()
    return table


def create_docx(markdown_text: str):
    doc = Document()
    setup_docx_page(doc)
    configure_docx_styles(doc)

    first_heading = True
    for item in parse_markdown_lines(markdown_text):
        kind = item[0]
        if kind == "line":
            line = item[1].rstrip()
            if not line:
                continue
            if line == "---":
                doc.add_page_break()
            elif line.startswith("# "):
                text = clean_inline_markdown(line[2:])
                paragraph = doc.add_heading(text, level=1)
                if first_heading:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(96)
                    paragraph.paragraph_format.space_after = Pt(24)
                    for run in paragraph.runs:
                        set_run_font(run, "Calibri")
                        run.font.size = Pt(22)
                        run.font.color.rgb = RGBColor.from_string("1F4D78")
                    first_heading = False
            elif line.startswith("## "):
                text = clean_inline_markdown(line[3:])
                if text.startswith("Ek - "):
                    doc.add_page_break()
                doc.add_heading(text, level=2)
            elif line.startswith("### "):
                doc.add_heading(clean_inline_markdown(line[4:]), level=3)
            elif line.startswith("- "):
                paragraph = add_docx_text(doc, line[2:], "List Bullet")
                style_screenshot_placeholder(paragraph)
            elif re.match(r"^\d+\.\s", line):
                add_docx_text(doc, re.sub(r"^\d+\.\s", "", line), "List Number")
            elif line.startswith("> "):
                add_docx_text(doc, line[2:], "Intense Quote")
            else:
                paragraph = add_docx_text(doc, line)
                style_screenshot_placeholder(paragraph)
        elif kind == "code":
            _, lang, code = item
            add_code_block(doc, code)
        elif kind == "table":
            _, header, rows = item
            add_formatted_table(doc, header, rows)

    doc.save(DOCX_OUT)


def register_fonts():
    arial = Path(r"C:\Windows\Fonts\arial.ttf")
    arial_bold = Path(r"C:\Windows\Fonts\arialbd.ttf")
    consolas = Path(r"C:\Windows\Fonts\consola.ttf")
    if arial.exists():
        pdfmetrics.registerFont(TTFont("Arial", str(arial)))
    if arial_bold.exists():
        pdfmetrics.registerFont(TTFont("Arial-Bold", str(arial_bold)))
    if consolas.exists():
        pdfmetrics.registerFont(TTFont("Consolas", str(consolas)))


def create_pdf(markdown_text: str):
    register_fonts()
    styles = getSampleStyleSheet()
    base_font = "Arial" if "Arial" in pdfmetrics.getRegisteredFontNames() else "Helvetica"
    bold_font = "Arial-Bold" if "Arial-Bold" in pdfmetrics.getRegisteredFontNames() else "Helvetica-Bold"
    code_font = "Consolas" if "Consolas" in pdfmetrics.getRegisteredFontNames() else base_font

    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName=base_font, fontSize=9, leading=12, spaceAfter=6)
    title = ParagraphStyle("Title", parent=body, fontName=bold_font, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=14)
    h1 = ParagraphStyle("H1", parent=body, fontName=bold_font, fontSize=15, leading=18, spaceBefore=10, spaceAfter=8, textColor=colors.HexColor("#1f4e79"))
    h2 = ParagraphStyle("H2", parent=body, fontName=bold_font, fontSize=12, leading=15, spaceBefore=8, spaceAfter=5, textColor=colors.HexColor("#244062"))
    h3 = ParagraphStyle("H3", parent=body, fontName=bold_font, fontSize=10.5, leading=13, spaceBefore=6, spaceAfter=4)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=12, bulletIndent=3)
    code_style = ParagraphStyle("Code", parent=body, fontName=code_font, fontSize=6.2, leading=7.4, leftIndent=6, rightIndent=4)

    story = []
    first_heading = True
    for item in parse_markdown_lines(markdown_text):
        kind = item[0]
        if kind == "line":
            line = item[1].rstrip()
            if not line:
                continue
            if line == "---":
                story.append(PageBreak())
            elif line.startswith("# "):
                text = html.escape(clean_inline_markdown(line[2:]))
                story.append(Paragraph(text, title if first_heading else h1))
                first_heading = False
            elif line.startswith("## "):
                story.append(Paragraph(html.escape(clean_inline_markdown(line[3:])), h2))
            elif line.startswith("### "):
                story.append(Paragraph(html.escape(clean_inline_markdown(line[4:])), h3))
            elif line.startswith("- "):
                story.append(Paragraph(html.escape(clean_inline_markdown(line[2:])), bullet, bulletText="•"))
            elif re.match(r"^\d+\.\s", line):
                story.append(Paragraph(html.escape(clean_inline_markdown(line)), body))
            elif line.startswith("> "):
                story.append(Paragraph(html.escape(clean_inline_markdown(line[2:])), body))
            else:
                story.append(Paragraph(html.escape(clean_inline_markdown(line)), body))
        elif kind == "code":
            _, lang, code = item
            story.append(Preformatted(code, code_style, maxLineLength=110))
            story.append(Spacer(1, 0.15 * cm))
        elif kind == "table":
            _, header, rows = item
            data = [[Paragraph(html.escape(clean_inline_markdown(c)), body) for c in header]]
            for row in rows:
                data.append([Paragraph(html.escape(clean_inline_markdown(c)), body) for c in row])
            table = Table(data, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9eaf7")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
                        ("FONTNAME", (0, 0), (-1, 0), bold_font),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9eabb8")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.2 * cm))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont(base_font, 8)
        canvas.setFillColor(colors.grey)
        canvas.drawRightString(A4[0] - 1.5 * cm, 0.9 * cm, f"Sayfa {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        rightMargin=1.35 * cm,
        leftMargin=1.35 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.25 * cm,
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main():
    REPORTS_DIR.mkdir(exist_ok=True)
    complete = build_complete_markdown()
    COMPLETE_MD.write_text(complete, encoding="utf-8")
    create_docx(complete)
    create_pdf(complete)
    print(f"Wrote {COMPLETE_MD}")
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {PDF_OUT}")


if __name__ == "__main__":
    main()
